import os, time, requests
from docx import Document

import glob

files = glob.glob('/Users/frei/Documents/watchedFolder/*')
for f in files:
    os.remove(f)

WATCHED_FOLDER = '/Users/frei/Documents/watchedFolder/'
EXPORTED_FOLDER = '/Users/frei/Desktop/'

def translate_file(file_path):
    if file_path.endswith('.docx'):
        document = Document(file_path)
        text = '\n'.join([para.text for para in document.paragraphs])
    elif file_path.endswith('.txt'):
        with open(file_path, 'r') as file:
            text = file.read()

    
    
    payload = {
        'q': text,
        'langpair': 'nob|nno'
    }
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.post('https://www.apertium.org/apy/translate', data=payload, headers=headers)
    translated_text = response.text.strip()

    print(translated_text)

    # new_document = Document()
    # new_document.add_paragraph(translated_text)
    # new_file_path = os.path.join(EXPORTED_FOLDER, os.path.basename(file_path))
    # new_document.save(new_file_path)

    # os.remove(file_path)

def watch_folder():
    while True:
        time.sleep(1)
        for file_name in os.listdir(WATCHED_FOLDER):
            file_path = os.path.join(WATCHED_FOLDER, file_name)
            if os.path.isfile(file_path) and (file_path.endswith('.docx') or file_path.endswith('.txt')):
                translate_file(file_path)
            else:
                print("Nothing yet")


if __name__ == '__main__':
    watch_folder()